import io
import json
import os
import re
import tempfile
import time
from typing import Dict, List, Optional
from urllib.parse import parse_qs, urlparse

import docx
import requests
import streamlit as st
from google import genai
from groq import Groq
from supabase import Client, create_client

# ------------------------------------------------------------------------------
# 1. КОНФИГУРАЦИЯ СТРАНИЦЫ И СТИЛИ
# ------------------------------------------------------------------------------
st.set_page_config(
    page_title="Lecture AI — Конспект & Проверка знаний",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown(
    """
    <style>
    /* Скрывает кнопку Deploy */
    [data-testid="stAppDeployButton"], .stDeployButton {
        display: none !important;
    }
    .stButton button {
        border-radius: 8px;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

GROQ_API_KEY = st.secrets.get("GROQ_API_KEY", "")
GEMINI_API_KEY = st.secrets.get("GEMINI_API_KEY", "")
SUPABASE_URL = st.secrets.get("SUPABASE_URL", "")
SUPABASE_KEY = st.secrets.get("SUPABASE_KEY", "")

MAX_FILE_SIZE_BYTES = 18 * 1024 * 1024

# Инициализация Supabase клиента
supabase: Optional[Client] = None
if SUPABASE_URL and SUPABASE_KEY:
    try:
        supabase = create_client(SUPABASE_URL, SUPABASE_KEY)
    except Exception as e:
        st.warning(f"⚠️ Не удалось подключиться к Supabase: {e}")

# ------------------------------------------------------------------------------
# 2. ИНИЦИАЛИЗА SESSION STATE
# ------------------------------------------------------------------------------
if "guest_mode" not in st.session_state:
    st.session_state.guest_mode = False

if "quiz_block1" not in st.session_state:
    st.session_state.quiz_block1 = []
if "quiz_block2" not in st.session_state:
    st.session_state.quiz_block2 = []
if "quiz_block3" not in st.session_state:
    st.session_state.quiz_block3 = []

if "current_block" not in st.session_state:
    st.session_state.current_block = 1
if "b1_idx" not in st.session_state:
    st.session_state.b1_idx = 0
if "b2_idx" not in st.session_state:
    st.session_state.b2_idx = 0
if "b3_idx" not in st.session_state:
    st.session_state.b3_idx = 0

if "total_score" not in st.session_state:
    st.session_state.total_score = 0

if "show_explanation" not in st.session_state:
    st.session_state.show_explanation = False
if "is_correct" not in st.session_state:
    st.session_state.is_correct = None

# ------------------------------------------------------------------------------
# 3. ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ И БАЗА ДАННЫХ
# ------------------------------------------------------------------------------
def extract_youtube_id(url: str) -> Optional[str]:
    """Извлекает ID видео из любых ссылок YouTube"""
    parsed = urlparse(url)
    if parsed.hostname in ("www.youtube.com", "youtube.com"):
        if parsed.path == "/watch":
            return parse_qs(parsed.query).get("v", [None])[0]
        if parsed.path.startswith(("/embed/", "/v/")):
            return parsed.path.split("/")[2]
    elif parsed.hostname == "youtu.be":
        return parsed.path.lstrip("/")
    return None


def get_youtube_transcript_or_audio(video_url: str, processor: "LectureProcessor") -> str:
    """Получает субтитры через надежный сторонний API Supadata"""
    api_key = st.secrets.get("SUPADATA_API_KEY", "")
    if not api_key:
        raise Exception("Не найден SUPADATA_API_KEY в настройках st.secrets.")

    url = f"https://api.supadata.ai/v1/youtube/transcript?url={video_url}"
    headers = {"x-api-key": api_key}

    response = requests.get(url, headers=headers)
    if response.status_code != 200:
        raise Exception(f"Ошибка сервиса: {response.text}")

    data = response.json()
    content = data.get("content", [])
    if not content:
        raise Exception("У этого видео нет доступных субтитров.")

    return " ".join([item.get("text", "") for item in content])


def call_gemini_with_retry(client, model, prompt, retries=5, delay=5):
    """Безопасный вызов Gemini API с повторными попытками при перегрузке (503)"""
    for i in range(retries):
        try:
            response = client.models.generate_content(model=model, contents=prompt)
            return response.text
        except Exception as e:
            # Если сервер перегружен (503), превышен лимит (429) или модель недоступна
            if any(err in str(e) for err in ["503", "UNAVAILABLE", "429"]):
                if i < retries - 1:
                    sleep_time = delay * (i + 1)  # Нарастающая пауза: 5с, 10с, 15с, 20с
                    time.sleep(sleep_time)
                    continue
            raise e

def create_docx_bytes(markdown_text: str) -> bytes:
    doc = docx.Document()
    lines = markdown_text.split("\n")
    for line in lines:
        clean_line = line.strip()
        if not clean_line:
            continue
        if clean_line.startswith("# "):
            doc.add_heading(clean_line.replace("# ", ""), level=1)
        elif clean_line.startswith("## "):
            doc.add_heading(clean_line.replace("## ", ""), level=2)
        elif clean_line.startswith("### "):
            doc.add_heading(clean_line.replace("### ", ""), level=3)
        elif clean_line.startswith("- ") or clean_line.startswith("* "):
            doc.add_paragraph(re.sub(r"\*\*|\*", "", clean_line[2:]), style="List Bullet")
        else:
            doc.add_paragraph(re.sub(r"\*\*|\*", "", clean_line))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def save_lecture_to_db(user_email: str, title: str, summary_md: str, quiz_json: dict, raw_transcript: str, youtube_url: Optional[str] = None):
    """Сохраняет сгенерированную лекцию в базу данных Supabase"""
    if not supabase or user_email == "guest@guest.com":
        return

    try:
        data = {
            "user_email": user_email,
            "title": title,
            "summary_md": summary_md,
            "quiz_json": quiz_json,
            "raw_transcript": raw_transcript,
            "youtube_url": youtube_url
        }
        supabase.table("lectures").insert(data).execute()
        st.toast("💾 Конспект автоматически сохранен в вашем Личном Кабинете!", icon="✅")
    except Exception as e:
        st.error(f"Ошибка сохранения в базу данных: {e}")


def load_user_lectures(user_email: str):
    """Загружает список конспектов пользователя из Supabase"""
    if not supabase or user_email == "guest@guest.com":
        return []
    try:
        response = supabase.table("lectures").select("*").eq("user_email", user_email).order("created_at", desc=True).execute()
        return response.data
    except Exception as e:
        st.error(f"Ошибка загрузки истории: {e}")
        return []

# ------------------------------------------------------------------------------
# 4. ЛОГИКА ОБРАБОТКИ ЛЕКЦИИ (LECTURE PROCESSOR)
# ------------------------------------------------------------------------------
class LectureProcessor:
    def __init__(self, groq_key: str = GROQ_API_KEY, gemini_key: str = GEMINI_API_KEY):
        self.groq_client = Groq(api_key=groq_key)
        self.gemini_client = genai.Client(api_key=gemini_key)
        self.gemini_model = "gemini-2.5-flash"

    def _transcribe_file(self, file_path: str) -> str:
        with open(file_path, "rb") as audio_file:
            transcription = self.groq_client.audio.transcriptions.create(
                file=(os.path.basename(file_path), audio_file.read()),
                model="whisper-large-v3",
                response_format="text",
            )
        return str(transcription)

    def transcribe_audio(self, file_path: str) -> str:
        file_size = os.path.getsize(file_path)
        if file_size <= MAX_FILE_SIZE_BYTES:
            return self._transcribe_file(file_path)

        st.warning("⚠️ Файл больше 18 МБ. Разбиваем на части для отправки в Groq...")
        full_transcript = []
        chunk_size = MAX_FILE_SIZE_BYTES
        total_chunks = (file_size // chunk_size) + 1
        progress_bar = st.progress(0)

        with open(file_path, "rb") as f:
            idx = 0
            while True:
                chunk_data = f.read(chunk_size)
                if not chunk_data:
                    break

                with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as tmp_chunk:
                    tmp_chunk.write(chunk_data)
                    chunk_path = tmp_chunk.name

                try:
                    part_text = self._transcribe_file(chunk_path)
                    full_transcript.append(part_text)
                finally:
                    if os.path.exists(chunk_path):
                        os.remove(chunk_path)

                idx += 1
                progress_bar.progress(min(idx / total_chunks, 1.0))

        progress_bar.empty()
        return "\n".join(full_transcript)

    def generate_content_and_quiz(self, text_or_url: str, target_lang: str, is_youtube: bool = False):
        lang_instructions: Dict[str, str] = {
            "auto": "Определи язык лекции и составь весь материал СТРОГО на этом же языке (KK / RU / EN).",
            "kk": "Составь весь материал СТРОГО на казахском языке (Қазақ тілінде).",
            "ru": "Составь весь материал СТРОГО на русском языке.",
            "en": "Составь весь материал СТРОГО на английском языке (English).",
        }
        instruction = lang_instructions.get(target_lang, lang_instructions["auto"])
        source_prompt = f"Посмотри и проанализируй это видео на YouTube: {text_or_url}" if is_youtube else f"Текст лекции для анализа:\n{text_or_url}"

        prompt = f"""
Ты — высококлассный преподаватель и методист.
{source_prompt}

Твоя задача — составить наглядно структурированный, разложенный "по полочкам" конспект и проверку знаний из 3 блоков.

ЯЗЫКОВОЕ ТРЕБОВАНИЕ:
{instruction}

ТРЕБОВАНИЯ К КОНСПЕКТУ:
1. Максимально подробный и понятный для студента. Придумай в самом начале краткий емкий Заголовок лекции (Заголовок 1 класса #).
2. Используй таблицы сравнений, структурированные списки, выдели ключевые даты, термины и примеры из практики.
3. Оформи схематичный кластер (логическую связь основных тем лекции).

ТРЕБОВАНИЯ К 3 БЛОКАМ ПРОВЕРКИ ЗНАНИЙ (JSON):
- БЛОК 1: 10 подробных вопросов с 4 вариантами ответов (A, B, C, D) для ключевой проверки.
- БЛОК 2: Квиз из 5 вопросов с заполнением пропусков. Формулируй вопрос с пропуском "[ ... ]", например: "В 1917 году произошло [ ... ], которое изменило ход истории." и давай 4 варианта ответов.
- БЛОК 3: 5 вопросов формата "Верно или Неверно" (True / False).

ОБЯЗАТЕЛЬНАЯ СТРУКТУРА ОТВЕТА:
Сначала напиши подробный конспект в формате Markdown.
В самом конце напиши JSON строго в следующем формате:

===QUIZ_JSON_START===
{{
  "block1": [
    {{
      "question": "Текст вопроса 1?",
      "options": ["Вариант A", "Вариант B", "Вариант C", "Вариант D"],
      "correct_index": 0,
      "explanation": "Подробное пояснение ответов."
    }}
  ],
  "block2": [
    {{
      "question": "Вопрос квиза 1?",
      "options": ["Вариант 1", "Вариант 2", "Вариант 3", "Вариант 4"],
      "correct_index": 1,
      "explanation": "Пояснение квиза."
    }}
  ],
  "block3": [
    {{
      "statement": "Утверждение для проверки",
      "is_true": true,
      "explanation": "Почему это правда или ложь."
    }}
  ]
}}
===QUIZ_JSON_END===
"""
        quiz_json = {"block1": [], "block2": [], "block3": []}

        try:
            raw_text = call_gemini_with_retry(
                client=self.gemini_client,
                model=self.gemini_model,
                prompt=prompt,
                retries=4,
                delay=4
            )
        except Exception as e:
            raise RuntimeError(f"Не удалось получить ответ от Gemini API: {e}")

        summary_md = raw_text

        if raw_text and "===QUIZ_JSON_START===" in raw_text and "===QUIZ_JSON_END===" in raw_text:
            parts = raw_text.split("===QUIZ_JSON_START===")
            summary_md = parts[0].strip()
            json_str = parts[1].split("===QUIZ_JSON_END===")[0].strip()
            try:
                quiz_json = json.loads(json_str)
            except Exception as e:
                st.error(f"Ошибка парсинга викторины: {e}")

        return summary_md, quiz_json

# ------------------------------------------------------------------------------
# 5. ИНТЕРАКТИВНЫЙ ИГРОВОЙ МОДУЛЬ (QUIZ)
# ------------------------------------------------------------------------------
def render_quiz_game():
    st.subheader("🎯 Проверка знаний")

    b1 = st.session_state.quiz_block1
    b2 = st.session_state.quiz_block2
    b3 = st.session_state.quiz_block3

    if not b1 and not b2 and not b3:
        st.info("Сначала сгенерируйте конспект лекции.")
        return

    curr_block = st.session_state.current_block

    # ЭКРАН ФИНАЛА
    if curr_block > 3:
        st.balloons()
        max_score = len(b1) + len(b2) + len(b3)
        user_score = st.session_state.total_score
        perc = int((user_score / max_score) * 100) if max_score > 0 else 0

        st.markdown(f"""
        ### 🏆 Все 3 блока успешно пройдены!
        Ваш итоговый результат: **{user_score} из {max_score} баллов** (**{perc}%**).
        """)

        if perc >= 80:
            st.success("🌟 Потрясающе! Вы отлично усвоили лекционный материал!")
        elif perc >= 50:
            st.warning("👍 Хорошая работа! Но стоит еще раз прочитать некоторые фрагменты.")
        else:
            st.error("📚 Рекомендуем повторно перечитать конспект лекции.")

        if st.button("🔄 Пройти проверку знаний заново", type="primary"):
            st.session_state.current_block = 1
            st.session_state.b1_idx = 0
            st.session_state.b2_idx = 0
            st.session_state.b3_idx = 0
            st.session_state.total_score = 0
            st.session_state.show_explanation = False
            st.rerun()
        return

    # БЛОК 1: ОСНОВНОЙ ТЕСТ
    if curr_block == 1:
        st.info("📌 **Блок 1 из 3: Вопросы по материалу (10 вопросов)**")
        idx = st.session_state.b1_idx
        total_q = len(b1)

        if idx >= total_q:
            st.success("🎉 Блок 1 завершен! Переходим к Блоку 2 (Квиз)...")
            if st.button("Перейти к Блоку 2 ➡️", type="primary"):
                st.session_state.current_block = 2
                st.session_state.show_explanation = False
                st.rerun()
            return

        item = b1[idx]
        st.progress((idx) / total_q)
        st.caption(f"Вопрос {idx + 1} из {total_q} | Очки: {st.session_state.total_score}")

        st.markdown(f"#### ❓ {item['question']}")
        selected = st.radio(
            "Выберите вариант ответа:",
            options=item["options"],
            key=f"b1_q_{idx}",
            disabled=st.session_state.show_explanation,
        )

        if not st.session_state.show_explanation:
            if st.button("✅ Ответить", type="primary", key=f"b1_btn_{idx}"):
                selected_idx = item["options"].index(selected)
                is_correct = selected_idx == item["correct_index"]
                st.session_state.is_correct = is_correct
                st.session_state.show_explanation = True
                if is_correct:
                    st.session_state.total_score += 1
                st.rerun()
        else:
            if st.session_state.is_correct:
                st.success("🎉 **Верно!**")
            else:
                correct_text = item["options"][item["correct_index"]]
                st.error(f"❌ **Неверно.** Правильный ответ: **{correct_text}**")
            st.info(f"💡 **Пояснение:** {item['explanation']}")

            if st.button("Следующий вопрос ➡️", type="primary", key=f"b1_next_{idx}"):
                st.session_state.b1_idx += 1
                st.session_state.show_explanation = False
                st.rerun()

    # БЛОК 2: КВИЗ С ВЫПАДАЮЩИМИ СПИСКАМИ
    elif curr_block == 2:
        st.info("📌 **Блок 2 из 3: Квиз с заполнением пропусков (5 вопросов)**")
        if not b2:
            st.warning("Нет вопросов для Блока 2.")
            return

        with st.form(key="block2_form"):
            user_answers = []
            for i, q in enumerate(b2):
                st.markdown(f"**Вопрос {i + 1}:** {q['question']}")
                options = ["-- Нажмите, чтобы выбрать ответ --"] + q["options"]
                selected = st.selectbox(
                    label=f"Выберите ответ для вопроса №{i + 1}:",
                    options=options,
                    key=f"b2_select_{i}",
                    disabled=st.session_state.show_explanation
                )
                user_answers.append(selected)
                st.write("")

            submit_btn = st.form_submit_button("✅ Проверить ответы", type="primary", disabled=st.session_state.show_explanation)

        if submit_btn:
            if any(ans == "-- Нажмите, чтобы выбрать ответ --" for ans in user_answers):
                st.warning("⚠️ Пожалуйста, закройте все пробелы перед проверкой!")
            else:
                st.session_state.show_explanation = True
                score_for_b2 = 0
                for i, q in enumerate(b2):
                    correct_text = q["options"][q["correct_index"]]
                    if user_answers[i] == correct_text:
                        score_for_b2 += 1
                st.session_state.total_score += score_for_b2
                st.rerun()

        if st.session_state.show_explanation:
            st.markdown("### 📊 Результаты проверки Блока 2:")
            for i, q in enumerate(b2):
                correct_text = q["options"][q["correct_index"]]
                user_ans = user_answers[i] if i < len(user_answers) else ""
                if user_ans == correct_text:
                    st.success(f"**Вопрос {i + 1}: ✅ Верно!**\n\nВаш ответ: *{user_ans}*")
                else:
                    st.error(f"**Вопрос {i + 1}: ❌ Неверно.**\n\nВаш ответ: *{user_ans}*\n\nПравильный ответ: **{correct_text}**")
                st.info(f"💡 **Пояснение:** {q['explanation']}")
                st.markdown("---")

            if st.button("Перейти к Блоку 3 ➡️", type="primary"):
                st.session_state.current_block = 3
                st.session_state.show_explanation = False
                st.rerun()

    # БЛОК 3: ВЕРНО / НЕВЕРНО
    elif curr_block == 3:
        st.info("📌 **Блок 3 из 3: Верно или Неверно (True / False)**")
        idx = st.session_state.b3_idx
        total_q = len(b3)

        if idx >= total_q:
            st.success("🎉 Все блоки успешно пройдены!")
            if st.button("Завершить и узнать результаты 🏆", type="primary"):
                st.session_state.current_block = 4
                st.rerun()
            return

        item = b3[idx]
        st.progress((idx) / total_q)
        st.caption(f"Вопрос {idx + 1} из {total_q} | Очки: {st.session_state.total_score}")

        st.markdown(f"#### 📢 {item['statement']}")
        user_choice = st.radio(
            "Утверждение является истинным?",
            options=["Верно", "Неверно"],
            key=f"b3_radio_{idx}",
            disabled=st.session_state.show_explanation,
        )

        if not st.session_state.show_explanation:
            if st.button("✅ Ответить", type="primary", key=f"b3_ans_btn_{idx}"):
                choice_bool = user_choice == "Верно"
                is_correct = choice_bool == item["is_true"]
                st.session_state.is_correct = is_correct
                st.session_state.show_explanation = True
                if is_correct:
                    st.session_state.total_score += 1
                st.rerun()
        else:
            if st.session_state.is_correct:
                st.success("🎉 **Правильно!**")
            else:
                correct_ans_str = "Верно" if item["is_true"] else "Неверно"
                st.error(f"❌ **Неверно.** Правильно: **{correct_ans_str}**")

            st.info(f"💡 **Пояснение:** {item['explanation']}")

            if st.button("Дальше ➡️", type="primary", key=f"b3_next_btn_{idx}"):
                st.session_state.b3_idx += 1
                st.session_state.show_explanation = False
                st.rerun()

# ------------------------------------------------------------------------------
# 6. ОКТРАН «ЛИЧНЫЙ КАБИНЕТ»
# ------------------------------------------------------------------------------
def render_dashboard(user_email: str):
    st.title("📂 Личный кабинет")
    st.subheader(f"История конспектов пользователя: `{user_email}`")

    if user_email == "guest@guest.com":
        st.warning("⚠️ Вы вошли в режиме Гостя. История конспектов не сохраняется в базе данных. Пожалуйста, авторизуйтесь через Google для сохранения лекций.")
        return

    lectures = load_user_lectures(user_email)

    if not lectures:
        st.info("У вас пока нет сохраненных конспектов. Создайте свой первый конспект на главной странице!")
        return

    for lec in lectures:
        with st.expander(f"📖 {lec['title']} (Создано: {lec['created_at'][:10]})"):
            st.markdown(lec["summary_md"])
            st.markdown("---")
            docx_bytes = create_docx_bytes(lec["summary_md"])
            st.download_button(
                label="📄 Скачать .docx",
                data=docx_bytes,
                file_name=f"{lec['title']}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"dl_{lec['id']}"
            )

# ------------------------------------------------------------------------------
# 7. ОСНОВНАЯ ЛОГИКА И ЭКРАН АВТОРИЗАЦИИ
# ------------------------------------------------------------------------------
def main():
# Проверка авторизации Streamlit (Google Auth)
    is_logged_in = False
    user_email = ""
    user_name = "Пользователь"

    # Безопасная проверка авторизации
    try:
        user_info = getattr(st, "user", None) or getattr(st, "experimental_user", None)
        if user_info and getattr(user_info, "is_logged_in", False):
            is_logged_in = True
            user_email = getattr(user_info, "email", "")
            user_name = getattr(user_info, "name", "") or user_email
    except AttributeError:
        is_logged_in = False

    if not is_logged_in and st.session_state.guest_mode:
        is_logged_in = True
        user_email = "guest@guest.com"
        user_name = "Гость"

        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            st.info("🔑 Авторизуйтесь, чтобы сохранять созданные конспекты в личном кабинете.")
            if st.button("🚀 Войти через Google", type="primary", use_container_width=True):
                st.login()

            st.markdown("<p style='text-align: center; color: gray;'>или</p>", unsafe_allow_html=True)

            if st.button("👤 Продолжить как Гость", use_container_width=True):
                st.session_state.guest_mode = True
                st.rerun()
        return

    # --- ВНУТРЕННИЙ ИНТЕРФЕЙС (Sidebar & Tabs) ---
    with st.sidebar:
        st.title("👤 Профиль")
        st.write(f"**Имя:** {user_name}")
        st.write(f"**Email:** {user_email}")

        if user_email != "guest@guest.com":
            if st.button("🚪 Выйти"):
                st.logout()
        else:
            if st.button("🔑 Войти через Google"):
                st.session_state.guest_mode = False
                st.login()

        st.markdown("---")
        st.header("⚙️ Настройки")
        selected_lang = st.selectbox(
            "Язык итогового конспекта",
            options=["auto", "kk", "ru", "en"],
            format_func=lambda x: {
                "auto": "🌐 Автоопределение",
                "kk": "🇰🇿 Қазақ тілі",
                "ru": "🇷🇺 Русский",
                "en": "🇬🇧 English",
            }[x],
        )

    # Переключение главных вкладок
    main_tab1, main_tab2 = st.tabs(["🚀 Создать конспект", "📂 Личный кабинет"])

    with main_tab2:
        render_dashboard(user_email)

    with main_tab1:
        st.title("🎓 Lecture AI — Генерация & Тестирование")
        st.caption("Автоматическое создание наглядных конспектов и интерактивной проверки знаний по аудио и видео лекциям.")

        # Выбор источника лекции
        source_type = st.radio(
            "Выберите способ загрузки лекции:",
            ("Загрузить аудиофайл", "Ссылка на YouTube"),
            horizontal=True,
        )

        processor = LectureProcessor()

        # Вариант 1: Загрузить аудиофайл
        if source_type == "Загрузить аудиофайл":
            uploaded_file = st.file_uploader(
                "Загрузите аудиозапись лекции (MP3, WAV, M4A, OGG)",
                type=["mp3", "wav", "m4a", "ogg"],
            )

            if uploaded_file and st.button("🚀 Начать обработку лекции", type="primary"):
                with tempfile.NamedTemporaryFile(
                    delete=False, suffix=os.path.splitext(uploaded_file.name)[1]
                ) as tmp_file:
                    tmp_file.write(uploaded_file.getbuffer())
                    temp_audio_path = tmp_file.name

                try:
                    with st.spinner("🎧 Расшифровка аудиозаписи (Groq Whisper)..."):
                        raw_transcript = processor.transcribe_audio(temp_audio_path)
                    
                    if not raw_transcript or not raw_transcript.strip():
                        st.error("❌ Не удалось распознать текст из аудиофайла.")
                        st.stop()

                    st.success("✅ Транскрибация завершена!")

                    with st.spinner("🤖 Gemini формирует структурированный конспект и 3 блока заданий..."):
                        res = processor.generate_content_and_quiz(
                            text_or_url=raw_transcript, target_lang=selected_lang, is_youtube=False
                        )
                        if not res:
                            raise ValueError("Модель не вернула ответ.")
                        summary_md, quiz_json = res

                    # Безопасная запись в session_state
                    st.session_state.summary_md = summary_md or ""
                    
                    if isinstance(quiz_json, dict):
                        st.session_state.quiz_block1 = quiz_json.get("block1", [])
                        st.session_state.quiz_block2 = quiz_json.get("block2", [])
                        st.session_state.quiz_block3 = quiz_json.get("block3", [])
                    else:
                        st.session_state.quiz_block1 = []
                        st.session_state.quiz_block2 = []
                        st.session_state.quiz_block3 = []

                    st.session_state.raw_transcript = raw_transcript
                    st.session_state.current_youtube_url = None

                    # Извлечение заголовка из markdown с защитой от None
                    lecture_title = uploaded_file.name
                    if summary_md and summary_md.strip():
                        first_line = summary_md.strip().split("\n")[0].replace("#", "").strip()
                        if first_line:
                            lecture_title = first_line

                    # Сохранение в Supabase
                    save_lecture_to_db(user_email, lecture_title, st.session_state.summary_md, quiz_json, raw_transcript)

                    # Сброс состояния игры
                    st.session_state.current_block = 1
                    st.session_state.b1_idx = 0
                    st.session_state.b2_idx = 0
                    st.session_state.b3_idx = 0
                    st.session_state.total_score = 0
                    st.session_state.show_explanation = False

                    # Перезапуск приложения для мгновенного отображения результатов
                    st.rerun()

                except Exception as e:
                    st.error(f"❌ Ошибка при обработке аудио: {str(e)}")
                finally:
                    if os.path.exists(temp_audio_path):
                        os.remove(temp_audio_path)
        # Вариант 2: Ссылка на YouTube
        else:
            youtube_url = st.text_input("Вставьте ссылку на видео с YouTube:")

            if youtube_url.strip() and st.button("🚀 Начать обработку YouTube видео", type="primary"):
                try:
                    with st.spinner("📜 Получение субтитров / аудио из видео..."):
                        transcript_text = get_youtube_transcript_or_audio(youtube_url, processor)

                    with st.spinner("🤖 Gemini анализирует текст и создает конспект и тесты..."):
                        summary_md, quiz_json = processor.generate_content_and_quiz(
                            text_or_url=transcript_text, target_lang=selected_lang, is_youtube=False
                        )

                    st.session_state.summary_md = summary_md
                    st.session_state.quiz_block1 = quiz_json.get("block1", [])
                    st.session_state.quiz_block2 = quiz_json.get("block2", [])
                    st.session_state.quiz_block3 = quiz_json.get("block3", [])
                    st.session_state.raw_transcript = transcript_text
                    st.session_state.current_youtube_url = youtube_url

                    # Извлечение заголовка из markdown
                    first_line = summary_md.split("\n")[0].replace("#", "").strip()
                    lecture_title = first_line if first_line else "YouTube Лекция"

                    # Сохранение в Supabase
                    save_lecture_to_db(user_email, lecture_title, summary_md, quiz_json, transcript_text, youtube_url)

                    # Сброс состояния игры
                    st.session_state.current_block = 1
                    st.session_state.b1_idx = 0
                    st.session_state.b2_idx = 0
                    st.session_state.b3_idx = 0
                    st.session_state.total_score = 0
                    st.session_state.show_explanation = False

                except Exception as e:
                    st.error(f"❌ Ошибка при обработке видео: {str(e)}")

        # Отображение результатов текущего сеанса
        if "summary_md" in st.session_state:
            st.markdown("---")
            tab_summary, tab_game, tab_transcript = st.tabs(
                ["📄 Подробный конспект", "🎯 Проверка знаний", "📜 Исходный транскрипт"]
            )

            with tab_summary:
                if st.session_state.get("current_youtube_url"):
                    st.video(st.session_state.current_youtube_url)
                st.markdown(st.session_state.summary_md)
                st.markdown("---")
                docx_file = create_docx_bytes(st.session_state.summary_md)
                st.download_button(
                    label="📄 Скачать конспект в формате Word (.docx)",
                    data=docx_file,
                    file_name="lecture_summary.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

            with tab_game:
                render_quiz_game()

            with tab_transcript:
                st.text_area("Распознанный исходный текст:", value=st.session_state.raw_transcript, height=350)


if __name__ == "__main__":
    main()
